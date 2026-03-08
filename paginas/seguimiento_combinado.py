"""Página Seguimiento Solicitud & Tareas — vista combinada por tarea."""

import io
import datetime
import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
import core.queries as q
from core.config import LABEL_ESTADO_SOL, LABEL_ESTADO_TRAB


def _fmt(val):
    if not val or val == "—":
        return "—"
    if hasattr(val, "strftime"):
        return val.strftime("%d-%m-%Y")
    if isinstance(val, str):
        try:
            return datetime.date.fromisoformat(val[:10]).strftime("%d-%m-%Y")
        except Exception:
            return val
    return str(val)


COLUMNAS = [
    "N° Solicitud", "Tarea", "Estado solicitud", "Estado tarea",
    "Fecha visita", "Fecha solicitud", "Fecha planificación", "Fecha ejecución",
]


def _generar_word(filas):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Márgenes carta
    for section in doc.sections:
        section.page_height = Cm(27.94)
        section.page_width  = Cm(21.59)
        section.left_margin = section.right_margin = Cm(2)
        section.top_margin  = section.bottom_margin = Cm(2)

    titulo = doc.add_paragraph()
    titulo.alignment = 1
    run = titulo.add_run("Seguimiento Solicitudes & Tareas")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(f"Generado: {datetime.date.today().strftime('%d-%m-%Y')}").runs[0].font.size = Pt(9)

    table = doc.add_table(rows=1, cols=len(COLUMNAS))
    table.style = "Table Grid"

    # Encabezado
    hdr = table.rows[0].cells
    for i, col in enumerate(COLUMNAS):
        hdr[i].text = col
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(8)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1A3A5C")
        shd.set(qn("w:color"), "FFFFFF")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Filas
    for fila in filas:
        row = table.add_row().cells
        for i, col in enumerate(COLUMNAS):
            row[i].text = str(fila.get(col) or "—")
            row[i].paragraphs[0].runs[0].font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _generar_pdf(filas):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    story  = []
    story.append(Paragraph("Seguimiento Solicitudes & Tareas", styles["Title"]))
    story.append(Paragraph(f"Generado: {datetime.date.today().strftime('%d-%m-%Y')}", styles["Normal"]))
    story.append(Spacer(1, 12))

    data = [COLUMNAS] + [[str(f.get(c) or "—") for c in COLUMNAS] for f in filas]

    col_widths = [2.5*cm, 5*cm, 2.5*cm, 2.5*cm, 2.2*cm, 2.2*cm, 2.2*cm, 2.2*cm]
    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND",  (0, 0), (-1, 0), colors.HexColor("#1A3A5C")),
        ("TEXTCOLOR",   (0, 0), (-1, 0), colors.white),
        ("FONTNAME",    (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",    (0, 0), (-1,-1), 7),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
        ("GRID",        (0, 0), (-1,-1), 0.4, colors.grey),
        ("ALIGN",       (0, 0), (-1,-1), "LEFT"),
        ("VALIGN",      (0, 0), (-1,-1), "MIDDLE"),
        ("TOPPADDING",  (0, 0), (-1,-1), 4),
        ("BOTTOMPADDING",(0,0), (-1,-1), 4),
    ]))
    story.append(t)
    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


def _html_preview(filas):
    filas_html = ""
    for i, f in enumerate(filas):
        bg = "#ffffff" if i % 2 == 0 else "#f0f4f8"
        celdas = "".join(f'<td style="padding:5px 8px;border:1px solid #ddd;font-size:11px;">{f.get(c) or "—"}</td>' for c in COLUMNAS)
        filas_html += f'<tr style="background:{bg}">{celdas}</tr>'

    encabezados = "".join(f'<th style="padding:7px 8px;background:#1a3a5c;color:white;font-size:11px;text-align:left;border:1px solid #1a3a5c;">{c}</th>' for c in COLUMNAS)

    return f"""
    <html><head><meta charset="utf-8">
    <style>
        body {{ font-family: Arial, sans-serif; margin: 2cm; }}
        h2 {{ color: #1a3a5c; font-size:15px; margin-bottom:4px; }}
        p  {{ font-size:10px; color:#666; margin:0 0 12px; }}
        table {{ width:100%; border-collapse:collapse; }}
        @media print {{ button {{ display:none; }} }}
    </style></head>
    <body>
    <h2>Seguimiento Solicitudes &amp; Tareas</h2>
    <p>Generado: {datetime.date.today().strftime('%d-%m-%Y')} — {len(filas)} registros</p>
    <table><thead><tr>{encabezados}</tr></thead><tbody>{filas_html}</tbody></table>
    </body></html>
    """


def render():
    st.title("Seguimiento Sol & Tareas")

    solicitudes = q.get_solicitudes()
    todas = [s for s in solicitudes if s["estado"] != "borrador"]

    if not todas:
        st.info("No hay solicitudes en seguimiento.")
        return

    # ── Construir filas ────────────────────────────────────────
    filas = []
    for sol in todas:
        trabajos = q.get_trabajos(sol["id"])
        activos  = [t for t in trabajos if t["estado"] != "cancelado"]
        for t in activos:
            filas.append({
                "_estado_sol":        sol["estado"],
                "_estado_tarea":      t["estado"],
                "N° Solicitud":       sol["numero"],
                "Tarea":              t["descripcion"],
                "Estado solicitud":   LABEL_ESTADO_SOL.get(sol["estado"], sol["estado"]),
                "Estado tarea":       LABEL_ESTADO_TRAB.get(t["estado"], t["estado"]),
                "Fecha visita":       _fmt(sol.get("fecha")),
                "Fecha solicitud":    _fmt(sol.get("created_at")),
                "Fecha planificación": _fmt(t.get("fecha_entrega")),
                "Fecha ejecución":    _fmt(t.get("fecha_termino")),
            })

    if not filas:
        st.info("No hay tareas para mostrar.")
        return

    # ── Filtros ────────────────────────────────────────────────
    estados_sol_raw   = sorted({f["_estado_sol"]   for f in filas})
    estados_tarea_raw = sorted({f["_estado_tarea"] for f in filas})
    lbl_sol   = ["Todos"] + [LABEL_ESTADO_SOL.get(e, e)  for e in estados_sol_raw]
    lbl_tarea = ["Todos"] + [LABEL_ESTADO_TRAB.get(e, e) for e in estados_tarea_raw]

    fc1, fc2 = st.columns(2)
    sel_sol   = fc1.selectbox("Filtrar x estado solicitud", lbl_sol,   key="sc_filtro_sol")
    sel_tarea = fc2.selectbox("Filtrar x estado tarea",     lbl_tarea, key="sc_filtro_tarea")

    filas_f = filas
    if sel_sol != "Todos":
        raw = next((e for e in estados_sol_raw   if LABEL_ESTADO_SOL.get(e, e)  == sel_sol),   None)
        filas_f = [f for f in filas_f if f["_estado_sol"] == raw]
    if sel_tarea != "Todos":
        raw = next((e for e in estados_tarea_raw if LABEL_ESTADO_TRAB.get(e, e) == sel_tarea), None)
        filas_f = [f for f in filas_f if f["_estado_tarea"] == raw]

    # ── Ordenar ────────────────────────────────────────────────
    orden_sol   = {"solicitada": 0, "planificada": 1, "en_ejecucion": 2, "completada": 3}
    orden_tarea = {"pendiente": 0, "planificado": 1, "en_ejecucion": 2, "completado": 3}

    import re
    def _num_sol(f):
        m = re.search(r"\d+", f["N° Solicitud"])
        return int(m.group()) if m else 0

    filas_f = sorted(
        filas_f,
        key=lambda f: (
            _num_sol(f),
            orden_sol.get(f["_estado_sol"], 99),
            orden_tarea.get(f["_estado_tarea"], 99),
        ),
    )

    df = pd.DataFrame(filas_f).drop(columns=["_estado_sol", "_estado_tarea"])

    st.dataframe(
        df,
        use_container_width=True,
        hide_index=True,
        column_config={
            "N° Solicitud":        st.column_config.TextColumn("N° Solicitud",        width="small"),
            "Tarea":               st.column_config.TextColumn("Tarea",               width="large"),
            "Estado solicitud":    st.column_config.TextColumn("Estado solicitud",    width="medium"),
            "Estado tarea":        st.column_config.TextColumn("Estado tarea",        width="medium"),
            "Fecha visita":        st.column_config.TextColumn("Fecha visita",        width="small"),
            "Fecha solicitud":     st.column_config.TextColumn("Fecha solicitud",     width="small"),
            "Fecha planificación": st.column_config.TextColumn("Fecha planificación", width="small"),
            "Fecha ejecución":     st.column_config.TextColumn("Fecha ejecución",     width="small"),
        },
    )

    # ── Botones Ver / Imprimir ─────────────────────────────────
    st.markdown("<div style='height:8px'></div>", unsafe_allow_html=True)
    col_ver, col_word, col_pdf = st.columns(3)

    if col_ver.button("Ver", use_container_width=True, key="btn_ver_sc"):
        st.session_state["sc_ver"] = not st.session_state.get("sc_ver", False)

    col_word.download_button(
        "Descargar Word",
        data=_generar_word(filas_f),
        file_name=f"seguimiento_{datetime.date.today().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
        key="btn_word_sc",
    )

    col_pdf.download_button(
        "Descargar PDF",
        data=_generar_pdf(filas_f),
        file_name=f"seguimiento_{datetime.date.today().strftime('%Y%m%d')}.pdf",
        mime="application/pdf",
        use_container_width=True,
        key="btn_pdf_sc",
    )

    # ── Vista previa ──────────────────────────────────────────
    if st.session_state.get("sc_ver"):
        st.divider()
        st.markdown("**Vista previa**")
        components.html(_html_preview(filas_f), height=500, scrolling=True)
